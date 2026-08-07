from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_docx(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                parts.append(" | ".join(values))
    return "\n\n".join(parts)
